"""
DOCX Generator Module
Generates Word documents with medical dialogue structure.
Supports styling for Doctor and Patient roles.
"""

import os
from datetime import datetime
from typing import Dict, List, Optional, Any
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from loguru import logger

logger.info("DOCX Generator module initialized")


class MedicalDialogueDocGenerator:
    """Генератор Word документов с медицинскими диалогами."""

    # Стили по умолчанию
    DEFAULT_STYLES = {
        'doctor': {
            'bold': True,
            'font_size': 11,
            'font_name': 'Arial',
            'color': RGBColor(0, 0, 100),  # Темно-синий
            'left_indent': None
        },
        'patient': {
            'bold': False,
            'font_size': 11,
            'font_name': 'Arial',
            'color': RGBColor(0, 0, 0),  # Черный
            'left_indent': Cm(1)  # Отступ 1 см
        },
        'title': {
            'font_size': 16,
            'bold': True,
            'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER
        },
        'subtitle': {
            'font_size': 12,
            'italic': True,
            'alignment': WD_PARAGRAPH_ALIGNMENT.CENTER
        },
        'metadata': {
            'font_size': 9,
            'italic': True,
            'color': RGBColor(100, 100, 100)
        },
        'timestamp': {
            'font_size': 9,
            'color': RGBColor(100, 100, 100)
        }
    }

    # Словарь маппинга спикеров на роли
    DEFAULT_SPEAKER_MAP = {
        'SPEAKER_00': 'Врач',
        'SPEAKER_01': 'Пациент',
        'SPEAKER_02': 'Спикер 3',
        'SPEAKER_03': 'Спикер 4',
        'DOCTOR': 'Врач',
        'PHYSICIAN': 'Врач',
        'PATIENT': 'Пациент',
        'CLINICIAN': 'Врач',
    }

    def __init__(
        self,
        styles: Dict = None,
        speaker_map: Dict = None,
        default_speaker: str = 'Спикер'
    ):
        """
        Инициализация генератора.

        Args:
            styles: Кастомные стили (опционально)
            speaker_map: Маппинг спикеров на роли (опционально)
            default_speaker: Название по умолчанию для неизвестных спикеров
        """
        self.styles = {**self.DEFAULT_STYLES, **(styles or {})}
        self.speaker_map = {**self.DEFAULT_SPEAKER_MAP, **(speaker_map or {})}
        self.default_speaker = default_speaker

        logger.debug("MedicalDialogueDocGenerator initialized")

    def generate(
        self,
        dialogue: List[Dict[str, Any]],
        metadata: Dict = None,
        include_timestamps: bool = True,
        include_speaker_labels: bool = True
    ) -> Document:
        """
        Сгенерировать Word документ с диалогом.

        Args:
            dialogue: Список сегментов диалога
            metadata: Метаданные (опционально)
            include_timestamps: Включать ли таймкоды
            include_speaker_labels: Включать ли метки спикеров

        Returns:
            Документ Word (docx.Document)
        """
        # Создание документа
        doc = Document()

        # Настройка стилей
        self._setup_styles(doc)

        # Заголовок
        self._add_title(doc, metadata)

        # Секция метаданных
        if metadata or include_timestamps:
            self._add_metadata(doc, metadata, include_timestamps)

        # Секция диалога
        self._add_dialogue(doc, dialogue, include_speaker_labels, include_timestamps)

        # Секция статистики
        self._add_statistics(doc, dialogue)

        logger.info(f"Документ сгенерирован: {len(dialogue)} сегментов")
        return doc

    def _setup_styles(self, doc: Document):
        """Настройка пользовательских стилей."""
        # Получаем стили
        style_dict = doc.styles

        # Настройка стиля для врача
        doctor_style = style_dict.add_style('Doctor', 1)
        doctor_style.base_style = style_dict['Normal']
        doctor_font = doctor_style.font
        doctor_font.name = self.styles['doctor']['font_name']
        doctor_font.size = Pt(self.styles['doctor']['font_size'])
        doctor_font.color.rgb = self.styles['doctor']['color']
        doctor_paragraph = doctor_style.paragraph_format
        doctor_paragraph.space_before = Pt(6)
        doctor_paragraph.space_after = Pt(6)

        # Настройка стиля для пациента
        patient_style = style_dict.add_style('Patient', 1)
        patient_style.base_style = style_dict['Normal']
        patient_font = patient_style.font
        patient_font.name = self.styles['patient']['font_name']
        patient_font.size = Pt(self.styles['patient']['font_size'])
        patient_font.color.rgb = self.styles['patient']['color']
        patient_paragraph = patient_style.paragraph_format
        patient_paragraph.left_indent = self.styles['patient']['left_indent']
        patient_paragraph.space_before = Pt(6)
        patient_paragraph.space_after = Pt(6)

    def _add_title(self, doc: Document, metadata: Dict = None):
        """Добавить заголовок документа."""
        title = doc.add_paragraph()
        title_run = title.add_run("Медицинская консультация")
        title_run.font.size = Pt(self.styles['title']['font_size'])
        title_run.font.bold = self.styles['title']['bold']
        title.alignment = self.styles['title']['alignment']

        # Подзаголовок с датой
        if metadata and metadata.get('date'):
            date_str = metadata['date']
        else:
            date_str = datetime.now().strftime("%d.%m.%Y")

        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run(f"Запись от: {date_str}")
        subtitle_run.font.size = Pt(self.styles['subtitle']['font_size'])
        subtitle_run.italic = self.styles['subtitle']['italic']
        subtitle.alignment = self.styles['subtitle']['alignment']

    def _add_metadata(
        self,
        doc: Document,
        metadata: Dict = None,
        include_timestamps: bool = True
    ):
        """Добавить метаданные."""
        meta_paragraph = doc.add_paragraph()

        meta_text = []
        if metadata:
            for key, value in metadata.items():
                if key != 'date':  # Дата уже в подзаголовке
                    meta_text.append(f"{key}: {value}")

        if include_timestamps:
            meta_text.append(f"Создано: {datetime.now().strftime('%H:%M:%S')}")

        meta_run = meta_paragraph.add_run(" | ".join(meta_text) if meta_text else "Медицинская консультация")
        meta_run.font.size = Pt(self.styles['metadata']['font_size'])
        meta_run.italic = self.styles['metadata']['italic']

    def _add_dialogue(
        self,
        doc: Document,
        dialogue: List[Dict[str, Any]],
        include_speaker_labels: bool = True,
        include_timestamps: bool = True
    ):
        """Добавить диалог в документ."""
        for i, segment in enumerate(dialogue):
            speaker = segment.get('speaker', 'Unknown')
            text = segment.get('text', '')
            start = segment.get('start')
            end = segment.get('end')

            # Получить роль спикера
            role = self._get_role(speaker)

            # Создать параграф
            p = doc.add_paragraph()

            # Добавить метку спикера
            if include_speaker_labels:
                role_run = p.add_run(f"{role}: ")
                role_run.font.bold = self.styles[role.lower()]['bold']
                role_run.font.size = Pt(self.styles[role.lower()]['font_size'])
                role_run.font.color.rgb = self.styles[role.lower()]['color']

                if role == 'Пациент':
                    p.paragraph_format.left_indent = self.styles['patient']['left_indent']

            # Добавить текст
            text_run = p.add_run(text)
            text_run.font.size = Pt(self.styles[role.lower()]['font_size'])
            text_run.font.color.rgb = self.styles[role.lower()]['color']

            # Добавить таймкоды
            if include_timestamps and start is not None and end is not None:
                time_str = f"[{self._format_time(start)}-{self._format_time(end)}]"
                time_run = p.add_run(f" {time_str}")
                time_run.font.size = Pt(self.styles['timestamp']['font_size'])
                time_run.font.italic = True
                time_run.font.color.rgb = self.styles['timestamp']['color']

            # Добавить пустую строку после каждого сегмента
            if i < len(dialogue) - 1:
                doc.add_paragraph()

    def _add_statistics(self, doc: Document, dialogue: List[Dict[str, Any]]):
        """Добавить статистику диалога."""
        if not dialogue:
            return

        doc.add_paragraph()

        # Подсчет статистики
        total_duration = sum(s.get('end', 0) - s.get('start', 0) for s in dialogue)
        total_words = sum(len(s.get('text', '').split()) for s in dialogue)
        speakers = set(s.get('speaker', 'Unknown') for s in dialogue)

        stats_text = [
            f"Длительность: {self._format_duration(total_duration)}",
            f"Слов: {total_words}",
            f"Спикеров: {len(speakers)}"
        ]

        stats_paragraph = doc.add_paragraph()
        stats_run = stats_paragraph.add_run(" | ".join(stats_text))
        stats_run.font.size = Pt(10)
        stats_run.font.bold = True

    def _get_role(self, speaker: str) -> str:
        """Получить роль по идентификатору спикера."""
        return self.speaker_map.get(speaker, self.default_speaker)

    def _format_time(self, seconds: float) -> str:
        """Форматировать время в MM:SS."""
        if seconds is None:
            return ""
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"

    def _format_duration(self, seconds: float) -> str:
        """Форматировать длительность в MM:SS."""
        if seconds is None:
            return "00:00"
        minutes = int(seconds // 60)
        secs = int(seconds % 60)
        return f"{minutes:02d}:{secs:02d}"

    def save(
        self,
        doc: Document,
        output_path: str
    ) -> str:
        """
        Сохранить документ.

        Args:
            doc: Документ Word
            output_path: Путь к выходному файлу

        Returns:
            Путь к сохраненному файлу
        """
        # Убедиться, что путь имеет расширение .docx
        output_path = Path(output_path)
        if output_path.suffix.lower() != '.docx':
            output_path = output_path.with_suffix('.docx')

        # Создать директорию, если её нет
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)

        # Сохранить
        doc.save(str(output_path))
        logger.info(f"Документ сохранен: {output_path}")
        return str(output_path)


def generate_dialogue_docx(
    dialogue: List[Dict[str, Any]],
    output_path: str,
    metadata: Dict = None,
    include_timestamps: bool = True,
    include_speaker_labels: bool = True
) -> str:
    """
    Упрощенная функция генерации документа.

    Args:
        dialogue: Список сегментов диалога
        output_path: Путь к выходному файлу
        metadata: Метаданные
        include_timestamps: Включать ли таймкоды
        include_speaker_labels: Включать ли метки спикеров

    Returns:
        Путь к сохраненному файлу
    """
    generator = MedicalDialogueDocGenerator()
    doc = generator.generate(
        dialogue,
        metadata=metadata,
        include_timestamps=include_timestamps,
        include_speaker_labels=include_speaker_labels
    )
    return generator.save(doc, output_path)


if __name__ == "__main__":
    # Пример использования
    sample_dialogue = [
        {'speaker': 'SPEAKER_00', 'text': 'Добрый день, что вас беспокоит?', 'start': 0, 'end': 2.5},
        {'speaker': 'SPEAKER_01', 'text': 'Здравствуйте, у меня боль в горле и температура.', 'start': 2.5, 'end': 6.0},
        {'speaker': 'SPEAKER_00', 'text': 'Как долго это длится и есть ли другие симптомы?', 'start': 6.0, 'end': 9.0},
    ]

    output_file = os.path.join(os.path.dirname(__file__), '..', 'output', 'sample_dialogue.docx')
    os.makedirs(os.path.dirname(output_file), exist_ok=True)

    doc = generate_dialogue_docx(
        sample_dialogue,
        output_file,
        metadata={'title': 'Медицинская консультация', 'date': '13.04.2026'}
    )

    print(f"Документ создан: {doc}")
